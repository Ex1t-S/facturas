from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt, RGBColor
from docx.oxml.ns import qn

from generate_fmh_template_proposals import set_cell_border, set_cell_margins


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "templates" / "propuestas-fmh-2026" / "01-industrial-clasica" / "01-presupuesto-industrial-clasica.docx"
SOURCE_REMITO = ROOT / "templates" / "propuestas-fmh-2026" / "01-industrial-clasica" / "01-remito-industrial-clasica.docx"
TARGET_DIR = ROOT / "templates" / "propuestas-fmh-2026" / "01-industrial-clasica-verde"
TARGET = TARGET_DIR / "01-presupuesto-industrial-clasica-verde.docx"
TARGET_REMITO = TARGET_DIR / "01-remito-industrial-clasica-verde.docx"
LOGO = ROOT / "templates" / "propuestas-fmh-2026" / "assets" / "fmh-logo.jpeg"


COLORS = {
    "17324D": "176245",  # primary -> deep green
    "2F4858": "00A65A",  # secondary -> FMH green
    "C89B3C": "00A65A",  # accent -> FMH green
    "EEF2F5": "F1F3F4",  # pale -> light gray
    "17212B": "4B5563",  # ink -> gray text
    "5C6873": "4B5563",  # muted -> gray text
}


def replace_color_attributes(doc: Document) -> None:
    """Recolor runs and cell/paragraph XML while preserving the template layout."""
    for paragraph in list(doc.paragraphs) + [p for table in doc.tables for row in table.rows for cell in row.cells for p in cell.paragraphs]:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb:
                old = str(run.font.color.rgb).upper()
                if old in COLORS:
                    run.font.color.rgb = RGBColor.from_string(COLORS[old])

    # Paragraph rules (including the horizontal accent above each section) live
    # outside table cell properties, so recolor every document XML node too.
    for node in doc.element.body.iter():
        for attr in (qn("w:fill"), qn("w:color"), qn("w:val")):
            value = node.get(attr)
            if value and value.upper() in COLORS:
                node.set(attr, COLORS[value.upper()])


def add_logo_to_brand_cell(doc: Document) -> None:
    table = doc.tables[0]
    cell = table.cell(0, 0)
    # Preserve the white logo background so the original green/white mark remains crisp.
    cell._tc.get_or_add_tcPr()
    from generate_fmh_template_proposals import set_cell_shading

    set_cell_shading(cell, "FFFFFF")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run()
    run.add_picture(str(LOGO), width=Mm(30))

    caption = cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)
    caption_run = caption.add_run("FABRICACIÓN Y MONTAJE INDUSTRIAL")
    caption_run.font.name = "Arial"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption_run.font.size = Pt(6.8)
    caption_run.font.bold = True
    caption_run.font.color.rgb = RGBColor.from_string("176245")
    set_cell_margins(cell, top=45, start=80, bottom=45, end=80)


def apply_green_variant() -> None:
    if not LOGO.exists():
        raise FileNotFoundError(f"No se encontró el logo: {LOGO}")
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, TARGET)
    doc = Document(TARGET)
    replace_color_attributes(doc)
    add_logo_to_brand_cell(doc)
    # Make the top-right document label the dark brand green and keep the white type.
    top_right = doc.tables[0].cell(0, 1)
    from generate_fmh_template_proposals import set_cell_shading

    set_cell_shading(top_right, "176245")
    doc.core_properties.title = "Presupuesto FMH · Variante Verde"
    doc.core_properties.subject = "Plantilla FMH verde y gris con logo"
    doc.save(TARGET)

    readme = TARGET_DIR / "README.md"
    readme.write_text(
        "# Variante verde FMH\n\n"
        "Copia controlada del modelo 01 (presupuesto), con el logo FMH incorporado y paleta "
        "verde/gris: verde principal `#00A65A`, verde oscuro `#176245`, gris claro `#F1F3F4` "
        "y texto `#4B5563`. Se conserva el contenido, la estructura simétrica del total y el formato A4.\n",
        encoding="utf-8",
    )


def apply_green_delivery_variant() -> None:
    if not SOURCE_REMITO.exists():
        raise FileNotFoundError(f"No se encontró el remito base: {SOURCE_REMITO}")
    TARGET_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_REMITO, TARGET_REMITO)
    doc = Document(TARGET_REMITO)
    replace_color_attributes(doc)
    add_logo_to_brand_cell(doc)
    top_right = doc.tables[0].cell(0, 1)
    from generate_fmh_template_proposals import set_cell_shading

    set_cell_shading(top_right, "176245")
    doc.core_properties.title = "Remito FMH · Variante Verde"
    doc.core_properties.subject = "Plantilla FMH verde y gris con logo"
    doc.save(TARGET_REMITO)


if __name__ == "__main__":
    apply_green_variant()
    apply_green_delivery_variant()
    print(TARGET)
    print(TARGET_REMITO)
