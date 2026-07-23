from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import sys
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUTPUT_ROOT = Path(__file__).resolve().parents[1] / "templates" / "propuestas-fmh-2026"


@dataclass(frozen=True)
class Theme:
    number: str
    slug: str
    name: str
    concept: str
    primary: str
    secondary: str
    accent: str
    pale: str
    ink: str
    muted: str
    header_mode: str


THEMES = (
    Theme(
        "01",
        "industrial-clasica",
        "Industrial Clásica",
        "Seria, sobria y muy legible. La opción más equilibrada para uso diario.",
        "17324D",
        "2F4858",
        "C89B3C",
        "EEF2F5",
        "17212B",
        "5C6873",
        "split",
    ),
    Theme(
        "02",
        "ingenieria-azul",
        "Ingeniería Azul",
        "Más técnica y modular, pensada para proyectos, silos y montajes.",
        "0B3558",
        "155E75",
        "22B8CF",
        "EAF7FA",
        "102A43",
        "52606D",
        "band",
    ),
    Theme(
        "03",
        "minimal-tecnica",
        "Minimal Técnica",
        "Blanco y negro, máxima claridad y excelente impresión económica.",
        "111827",
        "374151",
        "9CA3AF",
        "F3F4F6",
        "111827",
        "6B7280",
        "minimal",
    ),
    Theme(
        "04",
        "seguridad-industrial",
        "Seguridad Industrial",
        "Carácter de taller: grafito y naranja con jerarquía visual fuerte.",
        "262A2E",
        "454B50",
        "F28C28",
        "FFF4E8",
        "202428",
        "666D72",
        "block",
    ),
    Theme(
        "05",
        "campo-y-acero",
        "Campo & Acero",
        "Industrial y cercano al agro, con verde profundo y tonos cálidos.",
        "234B3A",
        "496A5A",
        "C6A15B",
        "F2F5EF",
        "1E2D27",
        "617067",
        "band",
    ),
)


QUOTE_ITEMS = (
    ("1", "Techado de galpón con 16 metros", "1", "servicio", "$ 50.000", "$ 50.000"),
    ("2", "Limpieza de cabezales de noria", "1", "servicio", "$ 20.000", "$ 20.000"),
    ("3", "Silo 200 t — fabricación y montaje", "1", "unidad", "$ 1.850.000", "$ 1.850.000"),
)

DELIVERY_ITEMS = (
    ("1", "Reparación y mejora de batea", "1", "trabajo"),
    ("2", "Limpieza de cabezales de noria", "1", "trabajo"),
    ("3", "Revisión y ajuste de transmisión", "1", "trabajo"),
)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, spec in edges.items():
        node = tc_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), spec.get("val", "single"))
        node.set(qn("w:sz"), str(spec.get("sz", 4)))
        node.set(qn("w:color"), spec.get("color", "D1D5DB"))
        node.set(qn("w:space"), "0")


def remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                start={"val": "nil"},
                end={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def fixed_table(table, widths_mm: Iterable[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for index, width in enumerate(widths_mm):
            row.cells[index].width = Mm(width)


def set_run(run, text: str, *, size=9, bold=False, color="111827", italic=False, font="Arial") -> None:
    run.text = text
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def reset_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def cell_text(
    cell,
    text: str,
    *,
    size=9,
    bold=False,
    color="111827",
    align=WD_ALIGN_PARAGRAPH.LEFT,
    italic=False,
    font="Arial",
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    reset_paragraph(paragraph, align=align)
    set_run(paragraph.add_run(), text, size=size, bold=bold, color=color, italic=italic, font=font)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_rule(doc: Document, color: str, size=12, space_after=4) -> None:
    paragraph = doc.add_paragraph()
    reset_paragraph(paragraph, after=space_after)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_label(doc: Document, title: str, theme: Theme) -> None:
    table = doc.add_table(rows=1, cols=2)
    fixed_table(table, (4, 178))
    remove_table_borders(table)
    set_cell_shading(table.cell(0, 0), theme.accent)
    cell_text(table.cell(0, 0), "")
    cell_text(table.cell(0, 1), title.upper(), size=8.5, bold=True, color=theme.primary)
    set_cell_margins(table.cell(0, 1), top=60, start=110, bottom=55, end=80)


def configure_document(doc: Document, theme: Theme, title: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)
    section.left_margin = Mm(14)
    section.right_margin = Mm(14)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(theme.ink)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    doc.core_properties.title = title
    doc.core_properties.subject = "Propuesta visual FMH Gestión 2026"
    doc.core_properties.author = "FMH Gestión"
    doc.core_properties.keywords = "FMH, presupuesto, remito, A4"

    footer = section.footer
    paragraph = footer.paragraphs[0]
    reset_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run(
        paragraph.add_run(),
        f"PROPUESTA {theme.number} · {theme.name.upper()} · FMH GESTIÓN",
        size=7,
        color=theme.muted,
    )


def add_brand_header(doc: Document, theme: Theme, document_type: str, number: str) -> None:
    if theme.header_mode == "minimal":
        add_rule(doc, theme.primary, size=16, space_after=2)
        table = doc.add_table(rows=1, cols=2)
        fixed_table(table, (118, 64))
        remove_table_borders(table)
        left, right = table.rows[0].cells
        cell_text(left, "FMH", size=27, bold=True, color=theme.primary)
        p = left.add_paragraph()
        reset_paragraph(p)
        set_run(p.add_run(), "FABRICACIÓN Y MONTAJE INDUSTRIAL", size=7.5, bold=True, color=theme.muted)
        cell_text(right, document_type.upper(), size=16, bold=True, color=theme.primary, align=WD_ALIGN_PARAGRAPH.RIGHT)
        p = right.add_paragraph()
        reset_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_run(p.add_run(), f"N.º {number}", size=9, bold=True, color=theme.muted)
    else:
        table = doc.add_table(rows=1, cols=2)
        fixed_table(table, (116, 66))
        remove_table_borders(table)
        left, right = table.rows[0].cells
        set_cell_shading(left, theme.primary)
        set_cell_shading(right, theme.secondary if theme.header_mode != "block" else theme.accent)
        cell_text(left, "FMH", size=27, bold=True, color="FFFFFF")
        p = left.add_paragraph()
        reset_paragraph(p)
        set_run(p.add_run(), "FABRICACIÓN Y MONTAJE INDUSTRIAL", size=7.5, bold=True, color="FFFFFF")
        cell_text(
            right,
            document_type.upper(),
            size=14,
            bold=True,
            color="FFFFFF" if theme.header_mode != "block" else theme.primary,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        p = right.add_paragraph()
        reset_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_run(
            p.add_run(),
            f"N.º {number}",
            size=9,
            bold=True,
            color="FFFFFF" if theme.header_mode != "block" else theme.primary,
        )
        for cell in (left, right):
            set_cell_margins(cell, top=150, start=170, bottom=130, end=170)

    contact = doc.add_table(rows=1, cols=2)
    fixed_table(contact, (111, 71))
    remove_table_borders(contact)
    cell_text(
        contact.cell(0, 0),
        "Silos · Norias · Sinfines · Estructuras metálicas",
        size=7.5,
        bold=True,
        color=theme.secondary,
    )
    cell_text(
        contact.cell(0, 1),
        "Huanguelén · 02923 648947 · fmharroyo@gmail.com",
        size=7.2,
        color=theme.muted,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    add_rule(doc, theme.accent, size=12, space_after=4)


def add_customer_block(doc: Document, theme: Theme, document_type: str) -> None:
    add_section_label(doc, "Datos del documento", theme)
    table = doc.add_table(rows=3, cols=4)
    fixed_table(table, (26, 72, 27, 57))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = (
        ("CLIENTE", "La Emancipación S.A.", "FECHA", "23/07/2026"),
        ("CUIT", "—", "MONEDA", "ARS"),
        ("DOMICILIO", "Huanguelén, Buenos Aires", "VALIDEZ" if document_type == "Presupuesto" else "PROYECTO", "15 días" if document_type == "Presupuesto" else "Silo 200 t"),
    )
    for row_index, row_data in enumerate(labels):
        row = table.rows[row_index]
        row.height = Mm(8)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for column_index, value in enumerate(row_data):
            cell = row.cells[column_index]
            if column_index in (0, 2):
                set_cell_shading(cell, theme.pale)
                cell_text(cell, value, size=7.5, bold=True, color=theme.primary)
            else:
                cell_text(cell, value, size=8.5, bold=row_index == 0 and column_index == 1, color=theme.ink)
            set_cell_border(
                cell,
                bottom={"val": "single", "sz": 4, "color": "D6DCE1"},
            )


def add_quote_table(doc: Document, theme: Theme) -> None:
    add_section_label(doc, "Detalle y valores", theme)
    table = doc.add_table(rows=1, cols=6)
    fixed_table(table, (9, 75, 14, 20, 29, 35))
    headers = ("#", "DESCRIPCIÓN", "CANT.", "UNIDAD", "P. UNITARIO", "SUBTOTAL")
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], theme.primary)
        cell_text(
            table.rows[0].cells[index],
            header,
            size=7.2,
            bold=True,
            color="FFFFFF",
            align=WD_ALIGN_PARAGRAPH.RIGHT if index >= 2 else WD_ALIGN_PARAGRAPH.LEFT,
        )
    set_repeat_table_header(table.rows[0])
    for row_index, item in enumerate(QUOTE_ITEMS):
        cells = table.add_row().cells
        table.rows[-1].height = Mm(10)
        table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, theme.pale)
        for column_index, value in enumerate(item):
            cell_text(
                cells[column_index],
                value,
                size=8,
                bold=column_index == 1,
                color=theme.ink,
                align=WD_ALIGN_PARAGRAPH.RIGHT if column_index in (0, 2, 4, 5) else WD_ALIGN_PARAGRAPH.LEFT,
            )
            set_cell_border(cells[column_index], bottom={"val": "single", "sz": 3, "color": "D9DEE3"})


def add_delivery_table(doc: Document, theme: Theme) -> None:
    add_section_label(doc, "Detalle de trabajos", theme)
    table = doc.add_table(rows=1, cols=4)
    fixed_table(table, (11, 111, 22, 38))
    headers = ("#", "DESCRIPCIÓN", "CANT.", "UNIDAD")
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], theme.primary)
        cell_text(
            table.rows[0].cells[index],
            header,
            size=7.5,
            bold=True,
            color="FFFFFF",
            align=WD_ALIGN_PARAGRAPH.RIGHT if index == 2 else WD_ALIGN_PARAGRAPH.LEFT,
        )
    set_repeat_table_header(table.rows[0])
    for row_index, item in enumerate(DELIVERY_ITEMS):
        cells = table.add_row().cells
        table.rows[-1].height = Mm(11)
        table.rows[-1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, theme.pale)
        for column_index, value in enumerate(item):
            cell_text(
                cells[column_index],
                value,
                size=8.5,
                bold=column_index == 1,
                color=theme.ink,
                align=WD_ALIGN_PARAGRAPH.RIGHT if column_index in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT,
            )
            set_cell_border(cells[column_index], bottom={"val": "single", "sz": 3, "color": "D9DEE3"})


def add_totals(doc: Document, theme: Theme, compact: bool = False) -> None:
    if compact:
        # Keep the label and amount in one compact right-aligned block.  Both
        # cells share the same fill and rule on the total row, so the amount
        # cannot visually drift away from its label.
        table = doc.add_table(rows=3, cols=2)
        fixed_table(table, (62, 48))
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        remove_table_borders(table)
        values = (("SUBTOTAL", "$ 1.920.000"), ("IVA 21%", "$ 403.200"), ("TOTAL", "$ 2.323.200"))
        for index, (label, value) in enumerate(values):
            left, right = table.rows[index].cells
            is_total = index == 2
            label_color = theme.primary if is_total else theme.muted
            cell_text(left, label, size=8.5 if is_total else 8, bold=is_total, color=label_color, align=WD_ALIGN_PARAGRAPH.RIGHT)
            cell_text(right, value, size=11.5 if is_total else 9, bold=True, color=theme.primary, align=WD_ALIGN_PARAGRAPH.RIGHT)
            for cell in (left, right):
                set_cell_margins(cell, top=55, start=75, bottom=55, end=75)
            if is_total:
                for cell in (left, right):
                    set_cell_shading(cell, theme.pale)
                    set_cell_border(
                        cell,
                        top={"val": "single", "sz": 8, "color": theme.accent},
                        bottom={"val": "single", "sz": 8, "color": theme.accent},
                    )
            else:
                set_cell_border(right, bottom={"val": "single", "sz": 3, "color": "D9DEE3"})
        return

    table = doc.add_table(rows=3, cols=2)
    fixed_table(table, (139, 43))
    remove_table_borders(table)
    values = (("SUBTOTAL", "$ 1.920.000"), ("IVA 21%", "$ 403.200"), ("TOTAL", "$ 2.323.200"))
    for index, (label, value) in enumerate(values):
        left, right = table.rows[index].cells
        cell_text(left, label, size=8, bold=index == 2, color=theme.muted, align=WD_ALIGN_PARAGRAPH.RIGHT)
        cell_text(right, value, size=9 if index < 2 else 12, bold=True, color=theme.primary, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if index == 2:
            set_cell_shading(right, theme.pale)
            set_cell_border(right, top={"val": "single", "sz": 8, "color": theme.accent})


def add_quote_terms(doc: Document, theme: Theme) -> None:
    add_section_label(doc, "Condiciones comerciales", theme)
    table = doc.add_table(rows=2, cols=2)
    fixed_table(table, (91, 91))
    remove_table_borders(table)
    conditions = (
        ("FORMA DE PAGO", "50% de anticipo · saldo contra entrega"),
        ("PLAZO ESTIMADO", "A coordinar según disponibilidad de materiales"),
        ("ALCANCE", "Materiales y mano de obra indicados en el detalle"),
        ("OBSERVACIONES", "Precios expresados en ARS. IVA discriminado."),
    )
    for index, (label, value) in enumerate(conditions):
        row, column = divmod(index, 2)
        cell = table.cell(row, column)
        set_cell_shading(cell, theme.pale)
        cell_text(cell, label, size=7, bold=True, color=theme.primary)
        paragraph = cell.add_paragraph()
        reset_paragraph(paragraph)
        set_run(paragraph.add_run(), value, size=7.8, color=theme.ink)
        set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
        set_cell_border(cell, bottom={"val": "single", "sz": 4, "color": "D6DCE1"})


def add_delivery_observations(doc: Document, theme: Theme) -> None:
    add_section_label(doc, "Observaciones", theme)
    table = doc.add_table(rows=1, cols=1)
    fixed_table(table, (182,))
    set_cell_shading(table.cell(0, 0), theme.pale)
    cell_text(
        table.cell(0, 0),
        "Trabajo realizado según lo coordinado con el cliente. Sin observaciones pendientes al momento de la entrega.",
        size=8.2,
        color=theme.ink,
    )
    table.rows[0].height = Mm(18)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_signature_block(doc: Document, theme: Theme, quote=False) -> None:
    spacer = doc.add_paragraph()
    reset_paragraph(spacer, after=2)
    set_run(spacer.add_run(), "", size=4)
    table = doc.add_table(rows=2, cols=2)
    fixed_table(table, (86, 86))
    remove_table_borders(table)
    labels = ("FMH · Responsable", "Aceptación del cliente" if quote else "Recibí conforme · Cliente")
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        cell_text(cell, "\n\n", size=8)
        set_cell_border(cell, bottom={"val": "single", "sz": 5, "color": theme.secondary})
        cell = table.cell(1, index)
        cell_text(cell, label, size=7.5, bold=True, color=theme.muted, align=WD_ALIGN_PARAGRAPH.CENTER)
    if not quote:
        p = doc.add_paragraph()
        reset_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=2)
        set_run(p.add_run(), "Aclaración · DNI · Fecha", size=7, color=theme.muted)


def add_disclaimer(doc: Document, theme: Theme, text: str) -> None:
    paragraph = doc.add_paragraph()
    reset_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=3)
    set_run(paragraph.add_run(), text, size=7, bold=True, color=theme.muted)


def build_quote(theme: Theme, target: Path) -> None:
    doc = Document()
    configure_document(doc, theme, f"FMH · Presupuesto · {theme.name}")
    add_brand_header(doc, theme, "Presupuesto", "00042")
    add_customer_block(doc, theme, "Presupuesto")
    add_quote_table(doc, theme)
    add_totals(doc, theme, compact=theme.number in {"01", "02", "03"})
    add_quote_terms(doc, theme)
    add_signature_block(doc, theme, quote=True)
    add_disclaimer(doc, theme, "Gracias por confiar en FMH · Esta propuesta no constituye factura.")
    doc.save(target)


def build_delivery_note(theme: Theme, target: Path) -> None:
    doc = Document()
    configure_document(doc, theme, f"FMH · Remito · {theme.name}")
    add_brand_header(doc, theme, "Remito", "00018")
    add_customer_block(doc, theme, "Remito")
    add_delivery_table(doc, theme)
    add_delivery_observations(doc, theme)
    add_signature_block(doc, theme, quote=False)
    add_disclaimer(doc, theme, "Documento no válido como factura.")
    doc.save(target)


def main() -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    manifest = []
    themes = THEMES[:3] if "--first-three" in sys.argv else THEMES
    for theme in themes:
        folder = OUTPUT_ROOT / f"{theme.number}-{theme.slug}"
        folder.mkdir(parents=True, exist_ok=True)
        quote_path = folder / f"{theme.number}-presupuesto-{theme.slug}.docx"
        delivery_path = folder / f"{theme.number}-remito-{theme.slug}.docx"
        build_quote(theme, quote_path)
        build_delivery_note(theme, delivery_path)
        manifest.append(
            {
                "number": theme.number,
                "name": theme.name,
                "concept": theme.concept,
                "quote": str(quote_path),
                "delivery_note": str(delivery_path),
            }
        )
    print({"output": str(OUTPUT_ROOT), "templates": manifest})


if __name__ == "__main__":
    main()
